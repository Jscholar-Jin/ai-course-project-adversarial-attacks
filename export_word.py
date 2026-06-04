import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".gif"}
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
BULLET_RE = re.compile(r"^-\s+(.*)$")
LINK_ONLY_RE = re.compile(r"^\s*(?:-\s+)?\[([^\]]+)\]\(([^)]+)\)\s*$")
INLINE_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")

FIGURE_TITLES = {
    "clean_acc": "Clean Accuracy Trade-off Across Models",
    "whitebox": "White-box Attack Strength at eps = 8/255",
    "white_eps": "White-box Attack Strength and Accuracy vs Epsilon",
    "pgd_steps": "PGD Gets Stronger with More Iterations",
    "transfer": "Transfer Attack from ResNet18 to CNN",
    "adv_train": "Adversarial Training Raises Robust Accuracy",
    "preprocess": "Preprocessing Defense Robust Accuracy",
    "adaptive": "Adaptive Attack Breaks Preprocessing Defenses",
    "vis_examples": "Adversarial Examples and Magnified Perturbations",
    "vis_tsne": "Feature Space Shift: Clean vs Adversarial Samples",
    "vis_fragile": "Most Fragile Classes Under Adversarial Attack",
}


def clean_text(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = INLINE_LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
    return text.strip()


def set_font(run, size=None, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3

    for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True


def image_path(root: Path, target: str) -> Path:
    path = Path(target)
    if not path.is_absolute():
        path = (root / path).resolve()
    return path


def add_paragraph(doc: Document, text: str, style: str = "Normal"):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(clean_text(text))
    set_font(run)
    return paragraph


def add_picture(doc: Document, path: Path):
    max_width = 6.2
    max_height = 8.1
    with Image.open(path) as img:
        width_px, height_px = img.size
    aspect = width_px / height_px
    width_in = max_width
    height_in = width_in / aspect
    if height_in > max_height:
        height_in = max_height
        width_in = height_in * aspect

    doc.add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    pic_paragraph = doc.paragraphs[-1]
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = FIGURE_TITLES.get(path.stem, path.stem.replace("_", " ").title())
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure. {caption}")
    set_font(run, size=10)
    run.italic = True


def export_markdown(markdown_path: Path, output_path: Path):
    root = markdown_path.parent
    doc = Document()
    configure_document(doc)

    in_code_block = False
    code_lines = []

    for raw_line in markdown_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                if code_lines:
                    block = doc.add_paragraph()
                    text = "\n".join(code_lines)
                    run = block.add_run(text)
                    set_font(run, size=10)
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        image_match = LINK_ONLY_RE.match(line)
        if image_match:
            label, target = image_match.groups()
            target_path = image_path(root, target)
            if target_path.suffix.lower() in IMAGE_EXTS and target_path.exists():
                add_picture(doc, target_path)
                continue
            add_paragraph(doc, f"{label}: {target}")
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            hashes, title = heading_match.groups()
            level = min(len(hashes), 3)
            paragraph = doc.add_heading(level=level)
            run = paragraph.add_run(clean_text(title))
            set_font(run, size={1: 18, 2: 15, 3: 13}[level], bold=True)
            continue

        ordered_match = ORDERED_RE.match(line)
        if ordered_match:
            paragraph = doc.add_paragraph(style="List Number")
            run = paragraph.add_run(clean_text(ordered_match.group(1)))
            set_font(run)
            continue

        bullet_match = BULLET_RE.match(line)
        if bullet_match:
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(clean_text(bullet_match.group(1)))
            set_font(run)
            continue

        add_paragraph(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default="report.md")
    parser.add_argument(
        "--output",
        default="CIFAR-10对抗攻击与防御实验报告.docx",
    )
    parser.add_argument("--copy-to", default="")
    return parser.parse_args()


def main():
    args = parse_args()
    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    export_markdown(input_path, output_path)

    if args.copy_to:
        target = Path(args.copy_to).resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(output_path, target)

    print(output_path)
    if args.copy_to:
        print(Path(args.copy_to).resolve())


if __name__ == "__main__":
    main()
